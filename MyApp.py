from win32com import client
import os
import os.path
import re
import docx
import mysql.connector
allname="小红|小明|小亮"
#全班人员姓名用|隔开需要输入
config = {
          'user':'root',
          'password':'root',
          'host':'192.168.74.128',
          'port':3306,
          'database':'test'}
conn = mysql.connector.connect(**config)
cursor = conn.cursor()
table=input("请输入表名")
col=input("请输入作业项目名")
sql = "SELECT name FROM "+table
cursor.execute(sql)
result_set = cursor.fetchall()
insert = "insert into "+table+" (name) values (%s)"
if result_set == []:#判断是否是空表如果是刚开始就是插入数据不是的话就是更新数据
    inname=allname.split('|')
    for x in inname:
        cursor.execute(insert,[x])
conn.commit()
addcol="ALTER TABLE "+table+"  ADD "+col+" INT(10)"
cursor.execute(addcol)
update="UPDATE "+table+" SET "+col+" = %s WHERE name = %s"
s=r"\d{2,3}"
#提取分数正则
sname=r"("+allname+")"
#班级人员姓名正则
pattern = re.compile(s)
namepattern=re.compile(sname)
rootdir=input("请输入文件夹绝对路径名")#每次只需改变目录位置
def doc2docx(doc_name, docx_name):
    try:
        # 首先将doc转换成docx
        word = client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_name)
        # 使用参数16表示将doc转换成docx
        doc.SaveAs(docx_name, 16)
        doc.Close()
        word.Quit()
        print(docx_name + "ok")
    except:
        pass
for parent, dirnames, filenames in os.walk(rootdir):  # 三个参数：分别返回1.父目录 2.所有文件夹名字（不含路径） 3.所有文件名字
    for filename in filenames:  # 输出文件信息
        # 提取文件后缀
        flag = re.split("\.", filename)
        if (flag[1] == "doc"):
            doc2docx(rootdir + "\\" + filename, rootdir + "\\" + flag[0] + ".docx")
            print(flag[0] + ".docx")
            os.remove(rootdir + "\\" + filename)#删除doc文件
j=0
for parent,dirnames,filenames in os.walk(rootdir):    #三个参数：分别返回1.父目录 2.所有文件夹名字（不含路径） 3.所有文件名字
    for filename in filenames:                        #输出文件信息
        # # 获取文档对象
        file = docx.Document(rootdir+"\\"+filename)
        # 输出段落编号及段落内容
        for i in range(len(file.paragraphs)):
            mark = file.paragraphs[i].text
            match = pattern.match(mark)
            if (match != None):
                tmp = str(match.group())#使用正则提取成绩
                mymark=tmp.split(":")[0]
                name = re.split(namepattern,filename)#使用正则从文件名中提取名字
                print(name[1]+"分数是" + mymark)#filename截取名字
                cursor.execute(update, [int(mymark), name[1]])
                j += 1
conn.commit()
cursor.close()
print(j)
# table.cell(0,0).value #单元格的值'
